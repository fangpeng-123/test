import os
from langchain_ollama import OllamaLLM
from langchain.chains.summarize import load_summarize_chain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# 加载文档函数
def load_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text)]
    
    elif ext == ".pdf":
        reader = PdfReader(file_path)
        texts = [page.extract_text() for page in reader.pages if page.extract_text()]
        return [Document(page_content=t) for t in texts]
    
    elif ext == ".docx":
        doc = DocxDocument(file_path)
        texts = [para.text for para in doc.paragraphs if para.text.strip()]
        return [Document(page_content=t) for t in texts]
    else:
        raise ValueError(f"不支持的文件格式：{ext}")
    
# 初始化大模型
ollm = OllamaLLM(model="qwen3:8b")

# 总结函数
def summarize_document(file_path):
    print("📄 正在加载文件...")
    docs = load_document(file_path)

    print("✂️ 正在切割文件...")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=200, chunk_overlap=10)
    split_docs = text_splitter.split_documents(docs)

    print("🧠 正在使用大模型进行总结...")
    chain = load_summarize_chain(ollm, chain_type="map_reduce", verbose=False)
    summary = chain.invoke(split_docs)

    return summary['output_text']

# 测试
if __name__ == "__main__":
    file_path = input("请输入你要总结的本地文件位置（支持txt/pdf/docx）：").strip()

    if not os.path.exists(file_path):
        print("❌ 文件不存在，请检查路径是否正确。")
    else:
        result = summarize_document(file_path)
        print("\n📝 最终总结结果：")
        print(result)